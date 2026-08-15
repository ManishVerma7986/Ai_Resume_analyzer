import os
import io
from typing import Tuple
from pypdf import PdfReader
import docx


def extract_text_from_pdf(file_path_or_buffer) -> Tuple[str, str]:
    """Extract text from a PDF file path or file-like buffer.

    Returns (text, error_message). error_message is empty on success.
    """
    try:
        if hasattr(file_path_or_buffer, "read"):
            reader = PdfReader(file_path_or_buffer)
        else:
            reader = PdfReader(open(file_path_or_buffer, "rb"))

        text_parts = []
        for page in reader.pages:
            try:
                page_text = page.extract_text() or ""
            except Exception:
                page_text = ""
            text_parts.append(page_text)

        text = "\n".join(text_parts).strip()
        if not text:
            return "", "No extractable text found (file may be scanned image-only PDF)."
        return text, ""
    except Exception as e:
        return "", f"PDF extraction error: {str(e)}"


def extract_text_from_docx(file_path_or_buffer) -> Tuple[str, str]:
    try:
        if hasattr(file_path_or_buffer, "read"):
            # python-docx expects a path or file-like; feed bytes
            file_stream = io.BytesIO(file_path_or_buffer.read())
            doc = docx.Document(file_stream)
        else:
            doc = docx.Document(file_path_or_buffer)

        paragraphs = [p.text for p in doc.paragraphs if p.text]
        text = "\n".join(paragraphs).strip()
        if not text:
            return "", "DOCX contains no extractable text."
        return text, ""
    except Exception as e:
        return "", f"DOCX extraction error: {str(e)}"


def extract_resume_text(uploaded_file) -> Tuple[str, str]:
    """Determine file type and extract text.

    uploaded_file may be a Streamlit UploadedFile or a filesystem path.
    Returns (text, error_message).
    """
    if uploaded_file is None:
        return "", "No file provided"

    filename = getattr(uploaded_file, "name", None) or str(uploaded_file)
    lower = filename.lower()
    try:
        if lower.endswith(".pdf"):
            return extract_text_from_pdf(uploaded_file)
        elif lower.endswith(".docx"):
            return extract_text_from_docx(uploaded_file)
        else:
            return "", "Unsupported file type. Only PDF and DOCX are supported."
    except Exception as e:
        return "", f"Extraction failed: {str(e)}"
